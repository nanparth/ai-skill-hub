"""DOCX adapter: parse paragraphs and tables into NormalizedDoc.

Heavy import (python-docx) is LAZY: imported inside parse(), never at module top.
"""
from __future__ import annotations
import re

from . import NormalizedDoc, DocBlock, DocTable, limit_value, mark_truncated

_HEADING_STYLE_RE = re.compile(r"heading\s*(\d+)", re.IGNORECASE)


def parse(src: str, **opts) -> NormalizedDoc:
    import docx  # lazy heavy import

    doc = NormalizedDoc(source_format="docx")
    document = docx.Document(src)
    max_paragraphs = limit_value(opts, "max_docx_paragraphs")
    max_tables = limit_value(opts, "max_docx_tables")
    max_table_rows = limit_value(opts, "max_docx_table_rows")

    current_heading = None
    heading_stack: list[tuple[int, str]] = []
    idx = 0
    processed_paragraphs = 0

    for para in document.paragraphs:
        text = (para.text or "").strip()
        if not text:
            idx += 1
            continue
        if max_paragraphs and processed_paragraphs >= max_paragraphs:
            mark_truncated(doc, "DOCX_PARAGRAPH_LIMIT_REACHED")
            break
        processed_paragraphs += 1

        style_name = ""
        if para.style is not None and para.style.name:
            style_name = para.style.name

        anchor = f"p{idx}"
        m = _HEADING_STYLE_RE.search(style_name)
        if m:
            level = int(m.group(1))
            current_heading = text
            # Trim stack to entries strictly above this level, then push.
            heading_stack = [(lvl, h) for lvl, h in heading_stack if lvl < level]
            heading_stack.append((level, text))
            doc.blocks.append(
                DocBlock(
                    text=text,
                    block_type="heading",
                    idx=idx,
                    level=level,
                    style=style_name,
                    anchor=anchor,
                    parent_heading=None,
                    heading_path=[h for _, h in heading_stack],
                )
            )
        elif "list" in style_name.lower() or _is_numbered(para):
            doc.blocks.append(
                DocBlock(
                    text=text,
                    block_type="list_item",
                    idx=idx,
                    style=style_name,
                    anchor=anchor,
                    parent_heading=current_heading,
                    heading_path=[h for _, h in heading_stack],
                )
            )
        else:
            doc.blocks.append(
                DocBlock(
                    text=text,
                    block_type="paragraph",
                    idx=idx,
                    style=style_name,
                    anchor=anchor,
                    parent_heading=current_heading,
                    heading_path=[h for _, h in heading_stack],
                )
            )
        idx += 1

    # Tables: row 0 = headers, remaining rows = data.
    table_rows_used = 0
    for t_no, table in enumerate(document.tables):
        if max_tables and t_no >= max_tables:
            mark_truncated(doc, "DOCX_TABLE_LIMIT_REACHED")
            break
        rows = []
        for row in table.rows:
            if max_table_rows and table_rows_used >= max_table_rows:
                mark_truncated(doc, "DOCX_TABLE_ROW_LIMIT_REACHED")
                break
            rows.append([(cell.text or "").strip() for cell in row.cells])
            table_rows_used += 1
        if not rows:
            continue
        headers = rows[0]
        data_rows = rows[1:]
        doc.tables.append(
            DocTable(
                headers=headers,
                rows=data_rows,
                anchor=f"table{t_no}",
                caption=current_heading,
            )
        )

    return doc


def _is_numbered(para) -> bool:
    """Detect a numbered/bulleted paragraph via its numbering properties."""
    try:
        pPr = para._p.pPr
        return pPr is not None and pPr.numPr is not None
    except AttributeError:
        return False
